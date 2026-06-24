# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
ROOT = OUT_DIR.parent
OUT_DOCX = OUT_DIR / "Ho_so_dang_ky_BCT_Dien_Tu_Hieu.docx"

COMPANY = {
    "name": "CÔNG TY TNHH MTV ĐIỆN MÁY HIẾU",
    "trade": "ĐIỆN MÁY HIẾU",
    "tax_code": "1402228630",
    "address": "166, Ấp Bình Thạnh 1, Xã Lấp Vò, Huyện Lấp Vò, Tỉnh Đồng Tháp, Việt Nam",
    "phone": "0979.553.289",
    "email": "Congvinh28@gmail.com",
    "representative": "TRẦN CÔNG VINH",
    "representative_title": "Giám đốc",
    "website_name": "Chợ Lấp Vò Online",
    "domain": "dienmayhieu.com",
    "website_url": "https://dienmayhieu.com",
    "hosting": "VinaHost",
    "issue_date": "06/04/2026",
    "issue_place": "Cơ quan đăng ký kinh doanh tỉnh Đồng Tháp",
    "employees": "1",
    "bct_user": "qltmdt@moit.gov.vn",
    "bct_api": "https://dienmayhieu.com/api_baocao_bct.php",
    "api_key_hash": "a40fb0f852327442863313008b13203cff842a7ca4a54ac68ecda0d6540f3b6e",
}


def get_or_add(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = get_or_add(rpr, "w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_style(style, size=12, bold=False, color=None, before=0, after=4, line=1.15):
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style._element.get_or_add_rPr()
    rfonts = get_or_add(rpr, "w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    styles = doc.styles
    configure_style(styles["Normal"], size=12, after=4, line=1.15)
    configure_style(styles["Heading 1"], size=14, bold=True, color="000000", before=12, after=6, line=1.15)
    configure_style(styles["Heading 2"], size=13, bold=True, color="000000", before=8, after=4, line=1.15)
    configure_style(styles["Heading 3"], size=12, bold=True, color="000000", before=6, after=3, line=1.15)
    configure_style(styles["List Bullet"], size=12, after=3, line=1.15)
    configure_style(styles["List Number"], size=12, after=3, line=1.15)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Hồ sơ đăng ký TMĐT - Điện Tử Hiếu | Trang ")
    set_run_font(run, size=10)
    add_page_field(footer_p)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, size=10)


def add_p(doc, text="", style=None, align=None, bold=False, italic=False, size=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_multirun_p(doc, parts, style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    for text, opts in parts:
        run = p.add_run(text)
        set_run_font(
            run,
            size=opts.get("size"),
            bold=opts.get("bold"),
            italic=opts.get("italic"),
            color=opts.get("color"),
        )
    return p


def set_cell_text(cell, text, bold=False, align=None, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, bold=bold, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = get_or_add(tc_pr, "w:shd")
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="B7B7B7", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_no_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = get_or_add(tc_pr, "w:tcMar")
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = get_or_add(tc_mar, f"w:{m}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_info_table(doc, rows, headers=None, widths=None):
    table = doc.add_table(rows=1 if headers else 0, cols=len(headers) if headers else len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if headers:
        for i, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(table.rows[0].cells[i], "F2F2F2")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, bold=False)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_margins(cell)
            if widths and i < len(widths):
                cell.width = Inches(widths[i])
    set_table_borders(table)
    return table


def add_label_table(doc, rows):
    return add_info_table(doc, rows, widths=[2.0, 4.1])


def add_official_header(doc, number=""):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = table.rows[0].cells
    left.width = Inches(2.7)
    right.width = Inches(3.6)
    set_cell_text(left, f"{COMPANY['name']}\nSố: {number or '.../ĐT-HS'}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    set_cell_text(
        right,
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11,
    )
    set_no_borders(table)
    add_p(doc, "Đồng Tháp, ngày 13 tháng 06 năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True)


def add_document_title(doc, title, subtitle=None):
    p = add_p(doc, title.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if subtitle:
        add_p(doc, subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)


def add_bullets(doc, items):
    for item in items:
        add_p(doc, item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        add_p(doc, item, style="List Number")


def add_signature(doc, recipient=True):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = table.rows[0].cells
    left.width = Inches(3.0)
    right.width = Inches(3.1)
    if recipient:
        set_cell_text(left, "Nơi nhận:\n- Như trên;\n- Lưu: Hồ sơ TMĐT.", size=11)
    else:
        set_cell_text(left, "", size=11)
    set_cell_text(
        right,
        "ĐẠI DIỆN CÔNG TY\nGIÁM ĐỐC\n(Ký tên, đóng dấu)\n\n\n\nTRẦN CÔNG VINH",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11,
    )
    set_no_borders(table)


def add_section_break(doc):
    doc.add_page_break()


def add_cover(doc):
    logo = ROOT / "logo.jpg"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo), width=Cm(2.4))

    add_p(doc, COMPANY["name"], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    add_p(doc, "MST: " + COMPANY["tax_code"], align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_p(doc, "", size=12)
    add_p(doc, "HỒ SƠ ĐĂNG KÝ WEBSITE CUNG CẤP DỊCH VỤ THƯƠNG MẠI ĐIỆN TỬ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    add_p(doc, COMPANY["website_name"], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    add_p(doc, COMPANY["website_url"], align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_p(doc, "", size=12)

    add_label_table(
        doc,
        [
            ("Tên miền chính", COMPANY["domain"]),
            ("Loại hình website", "Sàn giao dịch thương mại điện tử"),
            ("Dịch vụ chủ yếu", "Dịch vụ tuyển dụng lao động, việc làm; dịch vụ nội dung số và dịch vụ giá trị gia tăng"),
            ("Đơn vị hosting", COMPANY["hosting"]),
            ("Người đại diện pháp luật", f"{COMPANY['representative']} - {COMPANY['representative_title']}"),
            ("Số nhân viên quản lý website", COMPANY["employees"]),
            ("Đầu mối liên hệ", f"{COMPANY['phone']} - {COMPANY['email']}"),
        ],
    )
    add_p(doc, "Bộ hồ sơ gồm: (1) Đơn đăng ký; (2) Đề án cung cấp dịch vụ TMĐT; (3) Quy chế hoạt động website cung cấp dịch vụ TMĐT.", italic=True)


def add_registration_form(doc):
    add_official_header(doc, "01/2026/ĐK-TMĐT")
    add_document_title(doc, "Đơn đăng ký website cung cấp dịch vụ thương mại điện tử")
    add_p(doc, "Kính gửi: Cục Thương mại điện tử và Kinh tế số - Bộ Công Thương", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    add_p(doc, "1. Tên thương nhân, tổ chức sở hữu website cung cấp dịch vụ thương mại điện tử", style="Heading 2")
    add_bullets(doc, [f"Tên đăng ký: {COMPANY['name']}", f"Tên giao dịch: {COMPANY['trade']}", f"Mã số doanh nghiệp/Mã số thuế: {COMPANY['tax_code']}"])

    add_info_table(
        doc,
        [
            ("1", "Họ và tên người liên hệ", COMPANY["representative"], COMPANY["representative"]),
            ("2", "Chức danh", COMPANY["representative_title"], "Người chịu trách nhiệm website"),
            ("3", "Địa chỉ liên hệ", COMPANY["address"], COMPANY["address"]),
            ("4", "Số điện thoại di động", COMPANY["phone"], COMPANY["phone"]),
            ("5", "Địa chỉ email", COMPANY["email"], COMPANY["email"]),
            ("6", "Phương thức liên hệ trực tuyến", "Email, điện thoại, Zalo theo số điện thoại công ty", "Email, điện thoại, Zalo theo số điện thoại công ty"),
        ],
        headers=["STT", "Nội dung thông tin", "1a. Người đại diện theo pháp luật", "1b. Người chịu trách nhiệm website"],
        widths=[0.45, 1.45, 2.0, 2.0],
    )

    add_p(doc, "2. Địa chỉ trụ sở", style="Heading 2")
    add_p(doc, f"Địa chỉ: {COMPANY['address']}")
    add_p(doc, f"Điện thoại: {COMPANY['phone']}    Email: {COMPANY['email']}")

    add_p(doc, "3. Giấy chứng nhận đăng ký doanh nghiệp", style="Heading 2")
    add_p(
        doc,
        f"Loại giấy chứng nhận: Giấy chứng nhận đăng ký doanh nghiệp. Số: {COMPANY['tax_code']}. "
        f"Ngày cấp: {COMPANY['issue_date']}. Nơi cấp: {COMPANY['issue_place']}."
    )

    add_p(doc, "4. Tên miền Internet của website cung cấp dịch vụ TMĐT", style="Heading 2")
    add_bullets(
        doc,
        [
            f"Địa chỉ dành cho người mua/người sử dụng dịch vụ: {COMPANY['website_url']}",
            f"Địa chỉ dành cho người bán/nhà cung cấp/đối tác: {COMPANY['website_url']}",
            "Không đăng ký ứng dụng di động trong hồ sơ này.",
        ],
    )

    add_p(doc, "5. Loại hình dịch vụ cung cấp trên website", style="Heading 2")
    add_bullets(doc, ["[x] Dịch vụ sàn giao dịch thương mại điện tử", "[ ] Dịch vụ khuyến mại trực tuyến", "[ ] Dịch vụ đấu giá trực tuyến", "[ ] Dịch vụ khác"])

    add_p(doc, "6. Các loại hàng hóa hoặc dịch vụ chủ yếu được giao dịch trên website", style="Heading 2")
    add_bullets(
        doc,
        [
            "[x] Dịch vụ tuyển dụng lao động, việc làm, tư vấn du học",
            "[x] Dịch vụ nội dung số và dịch vụ giá trị gia tăng",
            "[x] Dịch vụ bảo dưỡng, sửa chữa, bảo trì, lắp đặt",
            "[x] Sản phẩm, dịch vụ khác: kết nối mua bán hàng hóa, dịch vụ địa phương, gọi thợ kỹ thuật và đối tác cung ứng dịch vụ.",
        ],
    )

    add_p(doc, "7. Đơn vị cung cấp dịch vụ hosting", style="Heading 2")
    add_p(doc, f"Đơn vị cung cấp dịch vụ hosting: {COMPANY['hosting']}.")

    add_p(doc, "8. Thông tin kết nối báo cáo với cơ quan quản lý", style="Heading 2")
    add_label_table(
        doc,
        [
            ("Phương án cung cấp báo cáo", "Qua API"),
            ("API báo cáo", COMPANY["bct_api"]),
            ("Tài khoản", COMPANY["bct_user"]),
            ("API_KEY_HASH đã cấu hình", COMPANY["api_key_hash"]),
            ("Cơ chế xác thực", "Gửi X-BCT-User và X-BCT-Secret qua header, hoặc username/password/api_key qua body. Hệ thống so khớp SHA-256 của API key với API_KEY_HASH."),
        ],
    )
    add_signature(doc)


def add_project_plan(doc):
    add_official_header(doc, "02/2026/ĐA-TMĐT")
    add_document_title(doc, "Đề án cung cấp dịch vụ thương mại điện tử", f"Website: {COMPANY['website_name']} - {COMPANY['domain']}")

    add_p(doc, "I. Giới thiệu tổng quan về thương nhân/tổ chức", style="Heading 1")
    add_p(doc, "1. Giới thiệu chung", style="Heading 2")
    add_p(
        doc,
        f"{COMPANY['name']} là công ty TNHH một thành viên do ông {COMPANY['representative']} làm người đại diện theo pháp luật, "
        f"mã số doanh nghiệp/mã số thuế {COMPANY['tax_code']}, ngày cấp {COMPANY['issue_date']}, trụ sở tại {COMPANY['address']}. "
        f"Công ty xây dựng và vận hành website {COMPANY['website_name']} tại tên miền {COMPANY['domain']} theo mô hình sàn giao dịch thương mại điện tử, "
        "kết nối người dùng với người bán, nhà cung cấp, đối tác dịch vụ và các nhóm dịch vụ số phù hợp."
    )

    add_p(doc, "Đầu mối tiếp nhận yêu cầu thông tin trực tuyến cho cơ quan quản lý nhà nước", style="Heading 3")
    add_info_table(
        doc,
        [
            ("1", "Họ và tên người liên hệ", COMPANY["representative"]),
            ("2", "Chức danh", COMPANY["representative_title"]),
            ("3", "Địa chỉ liên hệ", COMPANY["address"]),
            ("4", "Số điện thoại di động", COMPANY["phone"]),
            ("5", "Địa chỉ email", COMPANY["email"]),
            ("6", "Phương thức liên hệ trực tuyến", "Email, điện thoại, Zalo theo số điện thoại công ty"),
        ],
        headers=["STT", "Nội dung thông tin", "Thông tin đầu mối tiếp nhận"],
        widths=[0.5, 2.0, 3.6],
    )

    add_p(doc, "2. Nguồn vốn đầu tư cho website", style="Heading 2")
    add_bullets(
        doc,
        [
            "[x] Vốn doanh nghiệp",
            "[ ] Vốn ngân sách nhà nước",
            "[ ] Vốn đầu tư nước ngoài",
            "Mô tả: Website được đầu tư bằng nguồn vốn tự có của doanh nghiệp, triển khai theo từng giai đoạn phù hợp năng lực vận hành. Doanh nghiệp không có vốn đầu tư nước ngoài, không có nhân sự quản lý nước ngoài và không có hợp tác kỹ thuật bắt buộc với pháp nhân nước ngoài.",
        ],
    )

    add_p(doc, "3. Cơ cấu tổ chức và số lượng nhân viên", style="Heading 2")
    add_p(doc, f"Hiện tại số nhân sự trực tiếp quản lý, giám sát hoạt động website là {COMPANY['employees']} người. Giám đốc kiêm nhiệm quản trị vận hành, kiểm soát thông tin, tiếp nhận phản ánh và phối hợp xử lý yêu cầu của cơ quan quản lý.")
    add_info_table(
        doc,
        [
            ("Nhân viên kinh doanh", "0", "Kiêm nhiệm khi phát sinh", "Do Giám đốc trực tiếp phụ trách giai đoạn đầu"),
            ("Nhân viên IT", "0", "1", "Giám đốc kiêm quản trị kỹ thuật; nhà cung cấp hosting hỗ trợ hạ tầng"),
            ("Nhân viên pháp lý", "0", "1", "Giám đốc kiêm nhiệm, tham vấn khi cần"),
            ("Nhân viên hỗ trợ trực tuyến", "0", "1", "Tiếp nhận qua điện thoại/email"),
            ("Đội ngũ quản lý điều hành", "1", "0", "TRẦN CÔNG VINH - Giám đốc"),
        ],
        headers=["Bộ phận", "Chuyên trách", "Kiêm nhiệm", "Ghi chú"],
        widths=[2.0, 0.85, 0.85, 2.4],
    )

    add_p(doc, "4. Lĩnh vực hoạt động chủ yếu", style="Heading 2")
    add_bullets(
        doc,
        [
            "Sàn giao dịch thương mại điện tử kết nối người mua, người bán, nhà cung cấp và đối tác dịch vụ.",
            "Dịch vụ tuyển dụng lao động, việc làm và đăng tin nhu cầu nhân sự/dịch vụ tại địa phương.",
            "Dịch vụ nội dung số, dịch vụ giá trị gia tăng và các công cụ số hỗ trợ giao dịch.",
            "Dịch vụ bảo dưỡng, sửa chữa, bảo trì, lắp đặt và gọi thợ kỹ thuật.",
        ],
    )

    add_p(doc, "II. Mô hình tổ chức hoạt động cung cấp dịch vụ TMĐT", style="Heading 1")
    add_p(doc, "1. Loại hình cung cấp dịch vụ", style="Heading 2")
    add_bullets(doc, ["[x] Sàn giao dịch TMĐT", "[ ] Dịch vụ khuyến mại trực tuyến", "[ ] Dịch vụ đấu giá trực tuyến", "[ ] Dịch vụ khác"])
    add_p(
        doc,
        f"Website {COMPANY['website_name']} cho phép người dùng truy cập, tìm kiếm, gửi yêu cầu mua hàng hoặc sử dụng dịch vụ; "
        "cho phép người bán/nhà cung cấp/đối tác đăng thông tin, tiếp nhận yêu cầu, xác nhận giao dịch và phối hợp xử lý sau giao dịch. "
        "Ban quản lý website thực hiện kiểm soát thông tin, tiếp nhận phản ánh, lưu vết giao dịch và phối hợp với các bên liên quan."
    )
    add_bullets(
        doc,
        [
            f"Địa chỉ truy cập dành cho người mua/người sử dụng dịch vụ: {COMPANY['website_url']}",
            f"Địa chỉ truy cập dành cho người bán/nhà cung cấp/đối tác: {COMPANY['website_url']}",
            "Tài khoản thử nghiệm cho cơ quan quản lý sẽ được cung cấp qua kênh bảo mật khi có yêu cầu kiểm tra.",
        ],
    )

    add_p(doc, "2. Hoạt động cung cấp dịch vụ TMĐT", style="Heading 2")
    add_bullets(
        doc,
        [
            "Website dự kiến vận hành thử nghiệm từ tháng 06/2026 và hoàn thiện hồ sơ pháp lý trước khi mở rộng quy mô.",
            "Phạm vi hoạt động giai đoạn đầu: ưu tiên khu vực Lấp Vò, tỉnh Đồng Tháp; có thể mở rộng theo năng lực vận hành.",
            "Website có chức năng đặt hàng/gửi yêu cầu trực tuyến thông qua form đặt hàng, form liên hệ, thông báo trạng thái và kênh hỗ trợ.",
            "Các nhóm dịch vụ chủ yếu: việc làm, tuyển dụng; nội dung số và giá trị gia tăng; gọi thợ kỹ thuật; hàng hóa/dịch vụ địa phương phù hợp.",
        ],
    )

    add_p(doc, "3. Tiện ích và dịch vụ hỗ trợ", style="Heading 2")
    add_p(doc, "Điều kiện giao dịch chung", style="Heading 3")
    add_bullets(
        doc,
        [
            "Thông tin hàng hóa, dịch vụ, giá tham khảo, phí phát sinh và điều kiện phục vụ phải được thể hiện rõ ràng trước khi xác nhận giao dịch.",
            "Người dùng có quyền kiểm tra hàng hóa/dịch vụ trước khi thanh toán theo phương thức đã thỏa thuận.",
            "Chính sách đổi trả, hoàn tiền, bảo hành hoặc hỗ trợ sau bán hàng được xử lý theo tính chất từng giao dịch và quy định pháp luật liên quan.",
        ],
    )
    add_p(doc, "Phương thức thanh toán", style="Heading 3")
    add_bullets(
        doc,
        [
            "[x] Người mua - người bán tự thỏa thuận",
            "[x] Tiền mặt/COD",
            "[x] Chuyển khoản ngân hàng",
            "Website chưa lưu trữ thông tin thẻ thanh toán của khách hàng. Khi tích hợp thanh toán trực tuyến, công ty sẽ công bố nhà cung cấp trung gian thanh toán và cơ chế bảo mật tương ứng.",
        ],
    )
    add_p(doc, "Phương thức vận chuyển/giao nhận", style="Heading 3")
    add_bullets(
        doc,
        [
            "Người bán/nhà cung cấp tự vận chuyển hoặc tự thuê đơn vị vận chuyển.",
            "Đối với dịch vụ việc làm, nội dung số, gọi thợ hoặc dịch vụ tại chỗ, việc cung cấp dịch vụ được thực hiện theo lịch hẹn và địa điểm do các bên xác nhận.",
            "Trách nhiệm về chứng từ, nguồn gốc hàng hóa và điều kiện giao nhận thuộc về người bán/nhà cung cấp, trừ phần hỗ trợ lưu vết và điều phối của Ban quản lý website.",
        ],
    )
    add_p(doc, "Bảo vệ thông tin cá nhân", style="Heading 3")
    add_p(doc, "Công ty thu thập dữ liệu ở mức cần thiết để xác minh người dùng, xử lý đơn, điều phối đối tác, chăm sóc khách hàng, đối soát và tuân thủ yêu cầu quản lý. Dữ liệu được phân quyền truy cập, không bán cho bên thứ ba và chỉ chia sẻ ở mức cần thiết để thực hiện giao dịch hoặc theo yêu cầu hợp pháp của cơ quan nhà nước.")
    add_p(doc, "Quản lý và kiểm soát thông tin", style="Heading 3")
    add_bullets(
        doc,
        [
            "Ban quản lý kiểm tra thông tin người bán/nhà cung cấp/đối tác trước hoặc sau khi đăng tải tùy mức độ rủi ro.",
            "Thông tin cấm, sai sự thật, xâm phạm quyền sở hữu trí tuệ, hàng giả, hàng cấm, lừa đảo hoặc gây thiệt hại cho người tiêu dùng sẽ bị từ chối, ẩn, gỡ bỏ hoặc chuyển cơ quan có thẩm quyền.",
            "Công ty áp dụng danh sách từ khóa, rà soát thủ công và công cụ kỹ thuật phù hợp để hạn chế nội dung vi phạm.",
        ],
    )
    add_p(doc, "Giải quyết tranh chấp", style="Heading 3")
    add_p(doc, "Khi phát sinh khiếu nại, công ty tiếp nhận qua hotline/email, ghi nhận nội dung, yêu cầu các bên cung cấp chứng cứ, đối chiếu dữ liệu hệ thống và hỗ trợ hòa giải. Thời hạn xử lý ban đầu dự kiến trong 03 ngày làm việc kể từ khi nhận đủ thông tin cơ bản; vụ việc phức tạp được thông báo tiến độ cho các bên.")

    add_p(doc, "4. Nguồn thu và cơ chế thu phí", style="Heading 2")
    add_bullets(
        doc,
        [
            "Giai đoạn đầu ưu tiên hoàn thiện nền tảng, chưa áp dụng biểu phí cố định cho tất cả thành viên.",
            "Nguồn thu dự kiến có thể bao gồm phí dịch vụ, phí thành viên, quảng cáo, dịch vụ gia tăng hoặc tỷ lệ dựa trên giao dịch khi được công bố rõ trong chính sách của website.",
            "Doanh thu năm trước/năm nay: chưa phát sinh hoặc phát sinh theo giai đoạn vận hành thực tế, sẽ báo cáo theo dữ liệu kế toán và dữ liệu hệ thống.",
        ],
    )

    add_p(doc, "5. Biện pháp kỹ thuật - nghiệp vụ", style="Heading 2")
    add_bullets(
        doc,
        [
            f"Hệ thống website đặt tại hạ tầng hosting của {COMPANY['hosting']}, sử dụng tên miền {COMPANY['domain']} và kết nối HTTPS.",
            "Hệ thống quản trị phân quyền theo tài khoản, có cơ chế đăng nhập, lưu vết thao tác, chống truy cập trái phép ở mức ứng dụng.",
            "Dữ liệu giao dịch, đơn hàng, người dùng, đối tác và báo cáo được lưu trong cơ sở dữ liệu của hệ thống; công ty thực hiện sao lưu và kiểm tra định kỳ theo năng lực hạ tầng.",
            "Phương án kết nối báo cáo với Cổng quản lý hoạt động TMĐT: cung cấp API báo cáo tại " + COMPANY["bct_api"] + ".",
            "API hỗ trợ xác thực bằng tài khoản và API key; hệ thống lưu " + "BCT_REPORT_API_KEY_HASH" + " dạng SHA-256, không lưu API key thô trong mã nguồn.",
        ],
    )

    add_p(doc, "6. Hoạt động xúc tiến, tiếp thị dịch vụ", style="Heading 2")
    add_bullets(
        doc,
        [
            "Tiếp thị trực tuyến qua website, mạng xã hội, nội dung số và kênh hỗ trợ khách hàng.",
            "Tiếp thị trực tiếp tại địa phương, giới thiệu đến cửa hàng, thợ kỹ thuật, đối tác dịch vụ và khách hàng có nhu cầu.",
            "Các chương trình ưu đãi, tích điểm, QR hoặc quảng bá sẽ được công bố minh bạch trên website và tuân thủ quy định pháp luật.",
        ],
    )

    add_p(doc, "III. Phân định quyền và trách nhiệm giữa thương nhân/tổ chức với các bên sử dụng dịch vụ", style="Heading 1")
    add_p(doc, "1. Quyền và trách nhiệm của Ban quản lý website", style="Heading 2")
    add_bullets(
        doc,
        [
            "Cung cấp môi trường kỹ thuật để các bên đăng tin, đặt yêu cầu, xác nhận và lưu vết giao dịch.",
            "Ban hành, công bố và cập nhật quy chế hoạt động, chính sách bảo mật, chính sách xử lý khiếu nại và các điều kiện giao dịch liên quan.",
            "Kiểm soát, cảnh báo, ẩn/gỡ nội dung vi phạm; tạm khóa tài khoản hoặc quyền nhận việc khi phát hiện dấu hiệu gian lận, hàng cấm hoặc vi phạm pháp luật.",
            "Phối hợp với cơ quan nhà nước có thẩm quyền khi có yêu cầu cung cấp thông tin hợp pháp.",
        ],
    )
    add_p(doc, "2. Quyền và trách nhiệm của đối tác/người bán/nhà cung cấp", style="Heading 2")
    add_bullets(
        doc,
        [
            "Cung cấp thông tin đăng ký, hàng hóa, dịch vụ, giá, phí, điều kiện giao dịch và chứng từ liên quan một cách trung thực.",
            "Tự chịu trách nhiệm về chất lượng hàng hóa/dịch vụ, nguồn gốc, bảo hành, đổi trả, an toàn lao động và nghĩa vụ thuế/kế toán phát sinh.",
            "Không sử dụng dữ liệu khách hàng ngoài mục đích thực hiện giao dịch đã được xác nhận.",
        ],
    )
    add_p(doc, "3. Quyền và trách nhiệm của khách hàng/người sử dụng dịch vụ", style="Heading 2")
    add_bullets(
        doc,
        [
            "Cung cấp đúng thông tin liên hệ, địa chỉ, mô tả nhu cầu và phối hợp xác nhận giao dịch.",
            "Kiểm tra hàng hóa/dịch vụ, thanh toán theo thỏa thuận và phản ánh kịp thời khi có sự cố.",
            "Không tạo đơn ảo, lợi dụng khuyến mại, spam yêu cầu hoặc xâm phạm quyền, lợi ích hợp pháp của bên khác.",
        ],
    )
    add_p(doc, "4. Phân định trách nhiệm phát sinh từ giao dịch", style="Heading 2")
    add_p(doc, "Ban quản lý website chịu trách nhiệm về môi trường kỹ thuật, công cụ lưu vết, tiếp nhận phản ánh và hỗ trợ đối soát. Người bán/nhà cung cấp chịu trách nhiệm trực tiếp về hàng hóa, dịch vụ, chứng từ, giao nhận, bảo hành, đổi trả và các cam kết với khách hàng. Đơn vị vận chuyển hoặc đối tác thực hiện chịu trách nhiệm trong phạm vi công việc đã nhận.")
    add_signature(doc, recipient=False)


def add_regulation(doc):
    add_official_header(doc, "03/2026/QC-TMĐT")
    add_document_title(doc, "Quy chế hoạt động website cung cấp dịch vụ TMĐT", f"{COMPANY['website_name']} - {COMPANY['domain']}")

    add_p(doc, "I. Nguyên tắc chung", style="Heading 1")
    add_bullets(
        doc,
        [
            f"{COMPANY['website_name']} là nền tảng số do {COMPANY['name']} xây dựng và vận hành nhằm kết nối người mua, người bán, nhà cung cấp dịch vụ, đối tác việc làm, đối tác nội dung số và các nhóm dịch vụ phù hợp.",
            "Hoạt động của website tuân thủ quy định về thương mại điện tử, giao dịch điện tử, bảo vệ quyền lợi người tiêu dùng, bảo vệ dữ liệu cá nhân và các quy định pháp luật liên quan tại Việt Nam.",
            "Thông tin hàng hóa, dịch vụ, giá, phí, điều kiện cung cấp và trách nhiệm của từng bên phải được công bố trung thực, dễ kiểm tra.",
            f"Địa chỉ truy cập dành cho người mua/người sử dụng dịch vụ: {COMPANY['website_url']}.",
            f"Địa chỉ truy cập dành cho người bán/nhà cung cấp/đối tác: {COMPANY['website_url']}.",
        ],
    )

    add_p(doc, "II. Quy định chung", style="Heading 1")
    add_bullets(
        doc,
        [
            "Thành viên gồm người mua/người sử dụng dịch vụ, người bán/nhà cung cấp, đối tác việc làm, đối tác nội dung số, đối tác giao nhận/gọi thợ và Ban quản lý website.",
            "Thành viên phải cung cấp thông tin chính xác khi đăng ký, đặt hàng, đăng tin hoặc nhận yêu cầu. Ban quản lý có quyền yêu cầu bổ sung thông tin xác minh khi cần thiết.",
            "Không được đăng tải hàng hóa, dịch vụ, nội dung vi phạm pháp luật; không lừa đảo, gian lận, giả mạo, xâm phạm sở hữu trí tuệ, xâm phạm dữ liệu cá nhân hoặc gây thiệt hại cho bên khác.",
        ],
    )

    add_p(doc, "III. Quy trình giao dịch", style="Heading 1")
    add_p(doc, "1. Quy trình dành cho người mua/người sử dụng dịch vụ", style="Heading 2")
    add_numbers(
        doc,
        [
            "Truy cập website, chọn sản phẩm/dịch vụ hoặc nhập nhu cầu cần hỗ trợ.",
            "Cung cấp thông tin liên hệ, địa chỉ/vị trí, mô tả yêu cầu và thông tin cần thiết để xử lý.",
            "Kiểm tra thông tin, giá tham khảo, phí phát sinh nếu có và gửi yêu cầu/đơn hàng.",
            "Nhận xác nhận từ Ban quản lý, người bán, thợ, nhà cung cấp hoặc đối tác liên quan.",
            "Kiểm tra, nghiệm thu hàng hóa/dịch vụ và thanh toán theo phương thức đã thỏa thuận.",
        ],
    )
    add_p(doc, "2. Quy trình dành cho người bán/nhà cung cấp/đối tác", style="Heading 2")
    add_numbers(
        doc,
        [
            "Đăng ký hoặc cung cấp hồ sơ thông tin cho Ban quản lý website.",
            "Đăng thông tin hàng hóa, dịch vụ, năng lực cung cấp, giá, phí và điều kiện phục vụ.",
            "Tiếp nhận yêu cầu, xác nhận khả năng phục vụ, thời gian, chi phí và các điều kiện liên quan.",
            "Thực hiện giao dịch, cung cấp chứng từ, bảo hành hoặc hỗ trợ sau bán hàng theo cam kết.",
            "Phối hợp xử lý khiếu nại, đối soát và cập nhật trạng thái đơn hàng/dịch vụ.",
        ],
    )
    add_p(doc, "3. Quy trình giao nhận/vận chuyển hoặc cung cấp dịch vụ", style="Heading 2")
    add_p(doc, "Đối với hàng hóa vật lý, người bán tự vận chuyển hoặc thuê đơn vị vận chuyển phù hợp và chịu trách nhiệm về chứng từ, nguồn gốc, tình trạng hàng hóa cho đến khi bàn giao theo thỏa thuận. Đối với dịch vụ việc làm, nội dung số, gọi thợ hoặc dịch vụ tại chỗ, việc cung cấp được thực hiện theo lịch hẹn, phạm vi công việc và điều kiện do các bên xác nhận.")
    add_p(doc, "4. Quy trình xác nhận/hủy đơn hàng", style="Heading 2")
    add_p(doc, "Sau khi người dùng gửi yêu cầu, hệ thống ghi nhận và chuyển thông tin đến bên cung cấp phù hợp. Đơn hàng/yêu cầu được xác nhận khi bên cung cấp phản hồi khả năng phục vụ. Việc hủy đơn được chấp nhận khi các bên chưa thực hiện hoặc theo điều kiện đã công bố; trường hợp đã phát sinh chi phí thực tế, các bên thỏa thuận xử lý trên cơ sở chứng từ và dữ liệu hệ thống.")
    add_p(doc, "5. Quy trình đổi trả hàng và hoàn tiền", style="Heading 2")
    add_p(doc, "Người dùng gửi phản ánh qua hotline/email hoặc kênh hỗ trợ của website, cung cấp mã đơn, hình ảnh, nội dung lỗi và yêu cầu xử lý. Ban quản lý tiếp nhận, chuyển cho người bán/nhà cung cấp, đối chiếu dữ liệu và hỗ trợ phương án đổi trả, hoàn tiền hoặc sửa chữa theo tính chất giao dịch và quy định pháp luật.")
    add_p(doc, "6. Quy trình bảo hành/bảo trì", style="Heading 2")
    add_p(doc, "Đối với hàng hóa/dịch vụ có bảo hành, người bán/nhà cung cấp phải công bố điều kiện, thời hạn và phạm vi bảo hành. Website hỗ trợ lưu vết giao dịch và chuyển yêu cầu bảo hành đến bên có trách nhiệm xử lý.")
    add_p(doc, "7. Quy trình giải quyết tranh chấp, khiếu nại", style="Heading 2")
    add_numbers(
        doc,
        [
            "Người khiếu nại gửi thông tin qua hotline/email hoặc kênh hỗ trợ.",
            "Ban quản lý tiếp nhận, phân loại, yêu cầu bổ sung chứng cứ nếu cần.",
            "Các bên liên quan cung cấp giải trình, hình ảnh, chứng từ, tin nhắn hoặc dữ liệu liên quan.",
            "Ban quản lý hỗ trợ hòa giải trên cơ sở dữ liệu hệ thống và quy định đã công bố.",
            "Trường hợp không đạt thỏa thuận hoặc có dấu hiệu vi phạm pháp luật, vụ việc được chuyển cơ quan có thẩm quyền theo quy định.",
        ],
    )
    add_p(doc, "8. Quy trình khác", style="Heading 2")
    add_p(doc, "Các quy trình phát sinh mới sẽ được công bố, cập nhật trên website và áp dụng sau khi thông báo cho thành viên theo thời điểm phù hợp.")

    add_p(doc, "IV. Quy trình thanh toán", style="Heading 1")
    add_bullets(
        doc,
        [
            "Thanh toán giữa người mua và người bán/nhà cung cấp: tiền mặt/COD, chuyển khoản hoặc phương thức khác do các bên xác nhận.",
            "Thanh toán giữa đối tác và Ban quản lý website: thực hiện theo chính sách phí dịch vụ/thỏa thuận hợp tác được công bố hoặc ký kết riêng khi phát sinh.",
            "Website chưa lưu trữ thông tin thẻ thanh toán của khách hàng; khi có tích hợp thanh toán trực tuyến, công ty sẽ công bố nhà cung cấp và cơ chế bảo mật.",
        ],
    )

    add_p(doc, "V. Đảm bảo an toàn giao dịch", style="Heading 1")
    add_bullets(
        doc,
        [
            "Thông tin người bán/nhà cung cấp, hàng hóa/dịch vụ, giá, phí và điều kiện giao dịch được kiểm soát theo cơ chế của Ban quản lý.",
            "Thành viên được khuyến nghị chỉ thanh toán vào tài khoản hoặc phương thức đã được công bố/xác nhận.",
            "Website lưu vết đơn hàng, trạng thái xử lý, phản ánh và dữ liệu liên quan để hỗ trợ đối soát, khiếu nại và bảo vệ quyền lợi người dùng.",
            "Tài khoản vi phạm có thể bị cảnh báo, tạm khóa, gỡ nội dung, chấm dứt quyền sử dụng hoặc chuyển thông tin cho cơ quan có thẩm quyền.",
        ],
    )

    add_p(doc, "VI. Chính sách bảo vệ thông tin cá nhân của người tiêu dùng", style="Heading 1")
    add_p(doc, "Công ty thu thập dữ liệu cá nhân ở mức cần thiết gồm họ tên, số điện thoại, địa chỉ, vị trí khi người dùng cấp quyền, thông tin giao dịch, phản ánh, dữ liệu kỹ thuật và dữ liệu đối soát. Dữ liệu được dùng để xác minh, xử lý giao dịch, điều phối dịch vụ, chăm sóc khách hàng, phòng chống gian lận, cải thiện hệ thống và tuân thủ nghĩa vụ pháp lý.")
    add_p(doc, "Công ty không bán dữ liệu cá nhân cho bên thứ ba. Việc chia sẻ dữ liệu chỉ thực hiện ở mức cần thiết cho người bán, nhà cung cấp, đối tác giao nhận, kỹ thuật, thanh toán hoặc cơ quan nhà nước có thẩm quyền. Người dùng có thể yêu cầu kiểm tra, chỉnh sửa, hạn chế xử lý hoặc xóa dữ liệu trong phạm vi pháp luật cho phép.")

    add_p(doc, "VII. Quản lý thông tin trên website", style="Heading 1")
    add_bullets(
        doc,
        [
            "Ban quản lý có quyền rà soát, yêu cầu chỉnh sửa, từ chối hoặc gỡ bỏ thông tin vi phạm quy chế, vi phạm pháp luật hoặc gây rủi ro cho người tiêu dùng.",
            "Thông tin hàng hóa, dịch vụ phải mô tả đúng bản chất, nguồn gốc, giá, phí, điều kiện cung cấp và các hạn chế nếu có.",
            "Công ty áp dụng kiểm soát thủ công, tài khoản quản trị, nhật ký hệ thống, từ khóa cảnh báo và biện pháp kỹ thuật phù hợp để quản lý thông tin.",
        ],
    )

    add_p(doc, "VIII. Trách nhiệm trong trường hợp phát sinh lỗi kỹ thuật", style="Heading 1")
    add_p(doc, "Khi phát sinh lỗi kỹ thuật, lỗi đường truyền, phần mềm hoặc sự cố dữ liệu, Ban quản lý website tiếp nhận phản ánh, xác minh phạm vi ảnh hưởng, ưu tiên khắc phục, thông báo cho thành viên khi cần thiết và bảo vệ dữ liệu giao dịch trong khả năng kỹ thuật. Công ty không chịu trách nhiệm đối với lỗi do người dùng, thiết bị của người dùng, thông tin sai do thành viên cung cấp hoặc sự kiện bất khả kháng ngoài phạm vi kiểm soát hợp lý.")

    add_p(doc, "IX. Quyền và trách nhiệm của Ban quản lý website TMĐT", style="Heading 1")
    add_bullets(
        doc,
        [
            "Cung cấp, duy trì, nâng cấp nền tảng kỹ thuật và các công cụ hỗ trợ giao dịch.",
            "Công bố quy chế, chính sách, kênh liên hệ và tiếp nhận phản ánh.",
            "Kiểm soát nội dung, xử lý vi phạm, hỗ trợ giải quyết khiếu nại và phối hợp với cơ quan quản lý.",
            "Bảo vệ dữ liệu cá nhân, an toàn thông tin và lưu vết giao dịch theo năng lực hệ thống.",
        ],
    )

    add_p(doc, "X. Quyền và trách nhiệm của các bên tham gia website", style="Heading 1")
    add_bullets(
        doc,
        [
            "Người mua/người sử dụng dịch vụ có quyền được cung cấp thông tin rõ ràng, được hỗ trợ khiếu nại và được bảo vệ dữ liệu cá nhân; đồng thời phải cung cấp thông tin trung thực và thanh toán theo thỏa thuận.",
            "Người bán/nhà cung cấp/đối tác có quyền tiếp cận khách hàng phù hợp, sử dụng công cụ của website; đồng thời phải chịu trách nhiệm về hàng hóa, dịch vụ, chứng từ, chất lượng và nghĩa vụ pháp lý của mình.",
            "Các bên phải phối hợp thiện chí khi phát sinh tranh chấp, cung cấp chứng cứ trung thực và tuân thủ quyết định của cơ quan có thẩm quyền.",
        ],
    )

    add_p(doc, "XI. Điều khoản áp dụng", style="Heading 1")
    add_p(doc, "Quy chế này có hiệu lực kể từ ngày được công bố trên website và là căn cứ áp dụng cho hoạt động giao dịch trên nền tảng. Ban quản lý có quyền sửa đổi, bổ sung quy chế để phù hợp quy định pháp luật, yêu cầu quản lý, phạm vi dịch vụ và tình hình vận hành thực tế.")

    add_p(doc, "XII. Điều khoản cam kết", style="Heading 1")
    add_p(doc, f"{COMPANY['name']} cam kết vận hành website {COMPANY['website_name']} theo quy chế này, tuân thủ quy định pháp luật về thương mại điện tử, bảo vệ người tiêu dùng, bảo vệ dữ liệu cá nhân và nghĩa vụ báo cáo với cơ quan quản lý nhà nước có thẩm quyền.")
    add_label_table(
        doc,
        [
            ("Website", f"{COMPANY['website_name']} - {COMPANY['website_url']}"),
            ("Công ty/tổ chức", COMPANY["name"]),
            ("Địa chỉ", COMPANY["address"]),
            ("Tel/Hotline", COMPANY["phone"]),
            ("Email", COMPANY["email"]),
        ],
    )
    add_signature(doc, recipient=False)


def audit_document(doc):
    # Structural sanity checks before saving.
    assert len(doc.paragraphs) > 100
    assert len(doc.tables) >= 8
    for section in doc.sections:
        assert section.page_width.cm > 20
        assert section.page_height.cm > 29


def main():
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_section_break(doc)
    add_registration_form(doc)
    add_section_break(doc)
    add_project_plan(doc)
    add_section_break(doc)
    add_regulation(doc)
    audit_document(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
